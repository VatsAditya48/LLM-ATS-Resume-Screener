"""
AI Resume Analyser
------------------
Parses a job description and a folder of resumes (PDF/DOCX), scores each
candidate against the job description using an LLM (Groq), and prints a
ranked shortlist of top and bottom candidates.

Usage:
    python resume_analyser.py --resumes ./resumes --job-description ./job_description.txt
    python resume_analyser.py --resumes ./resumes --job-description ./job_description.txt --top 3 --bottom 3
    python resume_analyser.py --resumes ./resumes --job-description ./job_description.txt --output results.json
"""

import argparse
import json
import sys
import time
from pathlib import Path

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document


# ===========================
# Config
# ===========================

load_dotenv()
MODEL = "llama-3.3-70b-versatile"
SLEEP_BETWEEN_CALLS = 5  # seconds, to stay under rate limits


def get_client() -> Groq:
    import os

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError(
            "GROQ_API_KEY not found. Copy .env.example to .env and add your key."
        )
    return Groq(api_key=api_key)


# ===========================
# Pydantic Models
# ===========================

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]


class MatchResult(BaseModel):
    score: float
    details: dict


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


JOBD_SCHEMA = JobD.model_json_schema()
RESUME_SCHEMA = Resume.model_json_schema()


# ===========================
# LLM Calls
# ===========================

def parse_job_description(client: Groq, text: str) -> JobD:
    system_prompt = f"""
    You are an expert HR assistant.

    Your job is to analyze job descriptions and extract
    structured information from them.

    Return ONLY valid JSON matching this schema:

    {JOBD_SCHEMA}

    IMPORTANT:
    Do NOT return the schema itself.
    Do NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from the job description.

    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.
    """
    user_prompt = f"""
    Analyze the following job description:

    {text}
    """
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return JobD(**data)


def parse_resume(client: Groq, resume_text: str) -> Resume:
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {RESUME_SCHEMA}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return Resume(**data)


def final_score(client: Groq, job: JobD, resume: Resume) -> MatchResult:
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}

    Return JSON matching this schema:

    {match_schema}

    In the "details" field, include:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. A short final verdict

    "score" should be the overall match percentage from 0 to 100.
    Keep the response concise and easy to read.
    """
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)


# ===========================
# File Reading
# ===========================

def read_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file_path: Path) -> str:
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path: Path) -> str | None:
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


# ===========================
# CLI
# ===========================

def parse_args():
    parser = argparse.ArgumentParser(
        description="Score a folder of resumes against a job description using an LLM."
    )
    parser.add_argument(
        "--resumes",
        type=Path,
        default=Path("resumes"),
        help="Folder containing resume files (.pdf / .docx). Default: ./resumes",
    )
    parser.add_argument(
        "--job-description",
        type=Path,
        default=Path("job_description.txt"),
        help="Path to a text file containing the job description. "
             "Default: ./job_description.txt",
    )
    parser.add_argument(
        "--top",
        type=int,
        default=2,
        help="Number of top candidates to show (default: 2)",
    )
    parser.add_argument(
        "--bottom",
        type=int,
        default=2,
        help="Number of lowest-scoring candidates to show (default: 2)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Optional path to save full results as JSON.",
    )
    parser.add_argument(
        "--sleep",
        type=float,
        default=SLEEP_BETWEEN_CALLS,
        help=f"Seconds to sleep between LLM calls to avoid rate limits "
             f"(default: {SLEEP_BETWEEN_CALLS})",
    )
    return parser.parse_args()


def main():
    args = parse_args()

    if not args.job_description.exists():
        print(f"Job description file not found: {args.job_description}")
        print("Create it or pass --job-description path/to/file.txt")
        sys.exit(1)

    if not args.resumes.exists() or not args.resumes.is_dir():
        print(f"Resumes folder not found: {args.resumes}")
        print("Create it or pass --resumes path/to/folder")
        sys.exit(1)

    client = get_client()

    job_description_text = args.job_description.read_text(encoding="utf-8")
    job = parse_job_description(client, job_description_text)
    print(f"Job role detected: {job.role}")
    print(f"Minimum experience: {job.minimum_experience}")
    print(f"Education requirements: {job.education_requirements}\n")

    all_results = []
    resume_files = [
        f for f in args.resumes.iterdir()
        if f.suffix.lower() in [".pdf", ".docx"]
    ]

    if not resume_files:
        print(f"No .pdf or .docx files found in {args.resumes}")
        sys.exit(1)

    for file_path in resume_files:
        print(f"Processing: {file_path.name}")

        try:
            resume_text = read_resume(file_path)
            if not resume_text or not resume_text.strip():
                print(f"  Skipping {file_path.name}: no extractable text "
                      f"(likely a scanned/image file).")
                continue

            parsed_resume = parse_resume(client, resume_text)
            time.sleep(args.sleep)

            result = final_score(client, job, parsed_resume)
            time.sleep(args.sleep)

            print(f"  Score: {result.score}")
            all_results.append({
                "file": file_path.name,
                "name": parsed_resume.name or file_path.stem,
                "score": result.score,
                "details": result.details,
            })
        except Exception as e:
            print(f"  Failed to process {file_path.name}: {e}")
            continue

    if not all_results:
        print("\nNo candidates were successfully scored.")
        sys.exit(1)

    all_results.sort(key=lambda candidate: candidate["score"], reverse=True)

    # Guard against overlap when there are fewer candidates than requested
    n = len(all_results)
    top_n = min(args.top, n)
    bottom_n = min(args.bottom, max(0, n - top_n))

    top = all_results[:top_n]
    bottom = all_results[n - bottom_n:] if bottom_n else []

    print("\n" + "=" * 40)
    print("TOP CANDIDATES")
    print("=" * 40)
    for candidate in top:
        print(f"\n{candidate['name']} — {candidate['score']}%")
        print(json.dumps(candidate["details"], indent=2))

    if bottom:
        print("\n" + "=" * 40)
        print("LOWEST CANDIDATES")
        print("=" * 40)
        for candidate in bottom:
            print(f"\n{candidate['name']} — {candidate['score']}%")
            print(json.dumps(candidate["details"], indent=2))

    if args.output:
        args.output.write_text(json.dumps(all_results, indent=2), encoding="utf-8")
        print(f"\nFull results saved to {args.output}")


if __name__ == "__main__":
    main()
